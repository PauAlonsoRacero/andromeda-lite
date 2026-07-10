"""
Generador de documentos binarios reales (Word, Excel, PDF) a partir de texto.

Cuando la IA pide crear un archivo .docx/.xlsx/.pdf, no basta con escribir el
texto y ponerle esa extensión: un .txt renombrado a .docx no abre en Word. Aquí
convertimos el contenido de texto en el formato binario correcto.

Las dependencias (python-docx, openpyxl, reportlab) ya están en requirements.txt
y empaquetadas en el .exe/.app (hiddenimports del spec). Si por lo que sea no
están disponibles en tiempo de ejecución, caemos a texto plano y avisamos.
"""

from __future__ import annotations

from pathlib import Path

# Extensiones que requieren generación binaria.
BINARY_EXTS = {".docx", ".xlsx", ".pdf"}


def is_binary_doc(path: str) -> bool:
    return Path(path).suffix.lower() in BINARY_EXTS


def build_document(abs_path: Path, text: str) -> None:
    """Escribe `text` como documento binario según la extensión de `abs_path`.

    Lanza ImportError si falta la librería, para que el llamador caiga a texto.
    """
    ext = abs_path.suffix.lower()
    abs_path.parent.mkdir(parents=True, exist_ok=True)

    if ext == ".docx":
        _build_docx(abs_path, text)
    elif ext == ".xlsx":
        _build_xlsx(abs_path, text)
    elif ext == ".pdf":
        _build_pdf(abs_path, text)
    else:
        raise ValueError(f"Extensión no soportada para binario: {ext}")


def _build_docx(abs_path: Path, text: str) -> None:
    from docx import Document  # python-docx

    doc = Document()
    lines = text.splitlines() or [text]
    for line in lines:
        stripped = line.strip()
        # Encabezados markdown sencillos → estilos de Word
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        else:
            doc.add_paragraph(line)
    doc.save(str(abs_path))


def _build_xlsx(abs_path: Path, text: str) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    for line in text.splitlines():
        # Separar por comas o tabuladores si parece tabular; si no, una celda.
        if "\t" in line:
            cells = line.split("\t")
        elif "," in line:
            cells = [c.strip() for c in line.split(",")]
        else:
            cells = [line]
        ws.append(cells)
    wb.save(str(abs_path))


def _build_pdf(abs_path: Path, text: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(abs_path), pagesize=A4)
    width, height = A4
    x, y = 2 * cm, height - 2 * cm
    c.setFont("Helvetica", 11)
    for line in text.splitlines() or [text]:
        if y < 2 * cm:  # nueva página
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - 2 * cm
        # Recortar líneas muy largas para que quepan
        c.drawString(x, y, line[:110])
        y -= 0.6 * cm
    c.save()
