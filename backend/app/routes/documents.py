"""
documents.py — Generación de archivos reales (Word, Excel, PDF, Markdown).

Convierte el contenido de una respuesta (Markdown) en un documento descargable.
Esto es lo que permite a Andromeda "crear archivos": el modelo produce el
contenido y el backend lo materializa en un .docx/.xlsx/.pdf real.

POST /api/documents/generate
  body: { content: str, format: "docx"|"pdf"|"xlsx"|"md"|"txt", title?: str }
  → devuelve el archivo binario para descargar.
"""
from __future__ import annotations

import io
import logging
import re

from fastapi import APIRouter, Request
from fastapi.responses import JSONResponse, StreamingResponse

logger = logging.getLogger("andromeda.documents")
router = APIRouter()

MIME = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pdf":  "application/pdf",
    "md":   "text/markdown",
    "txt":  "text/plain",
}


@router.post("/generate", response_model=None)
async def generate_document(request: Request):
    body = await request.json()
    content = (body.get("content") or "").strip()
    fmt = (body.get("format") or "docx").lower()
    title = body.get("title") or "Documento Andromeda"

    if not content:
        return JSONResponse(status_code=400, content={"error": "Contenido vacío"})
    if fmt not in MIME:
        return JSONResponse(status_code=400, content={"error": f"Formato no soportado: {fmt}"})

    try:
        if fmt == "docx":
            data = _make_docx(content, title)
        elif fmt == "xlsx":
            data = _make_xlsx(content, title)
        elif fmt == "pdf":
            data = _make_pdf(content, title)
        else:  # md / txt
            data = content.encode("utf-8")
    except Exception as exc:
        logger.exception("Error generando documento")
        return JSONResponse(status_code=500, content={"error": f"No se pudo generar: {exc}"})

    safe_title = re.sub(r"[^\w\-]", "_", title)[:40] or "documento"
    filename = f"{safe_title}.{fmt}"
    return StreamingResponse(
        io.BytesIO(data),
        media_type=MIME[fmt],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Generadores por formato ───────────────────────────────────────────────────
def _make_docx(md: str, title: str) -> bytes:
    """Markdown → Word. Soporta encabezados, listas, negrita, código y tablas."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(title, level=0)

    in_code = False
    code_lines: list[str] = []
    table_buffer: list[str] = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            code_lines = []

    def flush_table():
        nonlocal table_buffer
        rows = [r for r in table_buffer if r.strip() and not re.match(r"^\s*\|[\s\-:|]+\|\s*$", r)]
        cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
        if cells:
            t = doc.add_table(rows=len(cells), cols=len(cells[0]))
            t.style = "Light Grid Accent 1"
            for i, row in enumerate(cells):
                for j, val in enumerate(row):
                    if j < len(t.rows[i].cells):
                        t.rows[i].cells[j].text = val
        table_buffer = []

    for line in md.split("\n"):
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        # Tablas markdown
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line)
            continue
        elif table_buffer:
            flush_table()

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif re.match(r"^\s*[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^\s*[-*]\s+", "", line), style="List Bullet")
        elif re.match(r"^\s*\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\s*\d+\.\s+", "", line), style="List Number")
        elif line.strip():
            p = doc.add_paragraph()
            _add_runs_with_bold(p, line)
        else:
            doc.add_paragraph()
    flush_code()
    flush_table()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_runs_with_bold(paragraph, text: str):
    """Renderiza **negrita** dentro de un párrafo de Word."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def _make_xlsx(md: str, title: str) -> bytes:
    """Extrae tablas markdown a una hoja de cálculo; si no hay, vuelca el texto."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31] or "Hoja1"

    rows = [r for r in md.split("\n") if r.strip().startswith("|")]
    if rows:
        for r in rows:
            if re.match(r"^\s*\|[\s\-:|]+\|\s*$", r):
                continue
            ws.append([c.strip() for c in r.strip().strip("|").split("|")])
    else:
        for line in md.split("\n"):
            ws.append([line])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_pdf(md: str, title: str) -> bytes:
    """Markdown → PDF legible con reportlab (encabezados, listas, código)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2*cm, rightMargin=2*cm)
    styles = getSampleStyleSheet()
    code_style = ParagraphStyle("code", parent=styles["Code"],
                                fontName="Courier", fontSize=8, leading=11,
                                backColor="#f4f4f4", alignment=TA_LEFT)
    flow = [Paragraph(_esc(title), styles["Title"]), Spacer(1, 12)]

    in_code = False
    code_lines: list[str] = []
    for line in md.split("\n"):
        if line.strip().startswith("```"):
            if in_code and code_lines:
                flow.append(Preformatted("\n".join(code_lines), code_style))
                flow.append(Spacer(1, 8))
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("### "):
            flow.append(Paragraph(_esc(line[4:]), styles["Heading3"]))
        elif line.startswith("## "):
            flow.append(Paragraph(_esc(line[3:]), styles["Heading2"]))
        elif line.startswith("# "):
            flow.append(Paragraph(_esc(line[2:]), styles["Heading1"]))
        elif re.match(r"^\s*[-*]\s+", line):
            txt = _esc(re.sub(r"^\s*[-*]\s+", "", line))
            flow.append(Paragraph(f"• {_bold(txt)}", styles["Normal"]))
        elif line.strip():
            flow.append(Paragraph(_bold(_esc(line)), styles["Normal"]))
            flow.append(Spacer(1, 4))
        else:
            flow.append(Spacer(1, 6))
    if code_lines:
        flow.append(Preformatted("\n".join(code_lines), code_style))

    doc.build(flow)
    return buf.getvalue()


def _esc(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _bold(text: str) -> str:
    # **negrita** → <b>negrita</b> para reportlab
    return re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
